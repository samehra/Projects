from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def report_generation_continuous(self):
    document = Document()
    document.add_heading('Data Analysis Report', 0)
    
    document.add_heading('Raw data summary', level=1)
    
    p1 = document.add_paragraph('The underlying data has ')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.add_run(str(len(self.y)))
    p1.add_run(' observations with ')
    if self.model_name != '9':
        p1.add_run(str(len(self.X[0])))
    else:
        p1.add_run('1')
    p1.add_run(' variables. The raw data was split into training and testing data using ')
    p1.add_run(self.report_partition)
    p1.add_run(', which resulted in a training data size of ')
    p1.add_run(str(len(self.y_train)))
    p1.add_run(' and a testing data size of ')
    p1.add_run(str(len(self.y_test)))
    p1.add_run('. ')
    p1.add_run(self.report_model)
    if (self.lamNum != 0):
        p1.add_run(' was used as the for data analysis and its tuning parameters were selected using ')
        p1.add_run(self.report_tuning)
    else:
        p1.add_run(' was used as the for data analysis')
    p1.add_run(' with the training dataset.')
    
    document.add_heading('Data analysis summary', level=1)
    
    p2 = document.add_paragraph('The goodness-of-fit measured in adj R-squared is ')
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.add_run(str(self.adj_rsquared))
    if self.model_name != '6':
        p2.add_run('. Based on the best model selected, the prediction error on the testing data is ')
        p2.add_run(str(self.rmse_test_final))
    else:
        p2.add_run('. Based on the best model selected, the prediction accuracy on the testing data is ')
        p2.add_run(str(self.testscore))
    p2.add_run('. The residual plot is attached in the following section')
    
    document.add_heading('The performance measure table', level=1)

    if self.model_name != '6':
        performance_measure_table = [ [1, "goodness-of-fit", self.adj_rsquared], [2, "training rmse", self.rmse_train_final]
        , [3, "testing rmse", self.rmse_test_final]]
    else:
        performance_measure_table = [[1, "goodness-of-fit", self.adj_rsquared],[2, "training accuracy", self.trainscore]
        , [3, "training type I error", self.T1Error],[4, "training type II error", self.T2Error],[5, "testing accuracy", self.testscore],[6, "training type I error", self.T1Error1],[7, "training type II error", self.T2Error1]]
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Performance measure index'
    hdr_cells[1].text = 'Performance measure name'
    hdr_cells[2].text = 'Value'
    for item in performance_measure_table:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item[0])
        row_cells[1].text = str(item[1])
        row_cells[2].text = str(item[2])
        
    document.add_heading('The residual plot', level=1)
    if self.model_name == '7':
        for i in range(self.y_train.shape[1]):
            document.add_heading('task '+str(i+1), level=2)
            document.add_picture('./'+self.folderName+'/residual_plot'+str(i)+'.png', height=Inches(4),width=Inches(5.5))
    else:
        document.add_picture('./' + self.folderName + '/residual_plot.png', height=Inches(4),width=Inches(5.5))
    if (self.lamNum != 0):
        document.add_heading('The tuning plot', level=1)
        document.add_picture('./'+self.folderName+'/parameter_tuning_plot.png', height=Inches(4),width=Inches(5.5))
        
    document.add_page_break()
    document.save('./'+self.folderName+'/report_demo.docx')
    
def report_generation_categorical(self):
    document = Document()
    document.add_heading('Data Analysis Report', 0)
    
    document.add_heading('Raw data summary', level=1)
    
    p1 = document.add_paragraph('The underlying data has ')
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.add_run(str(len(self.y)))
    p1.add_run(' observations with ')
    if self.model_name != '9':
        p1.add_run(str(len(self.X[0])))
    else:
        p1.add_run('1')
    p1.add_run(' variables. The raw data was split into training and testing data using ')
    p1.add_run(self.report_partition)
    p1.add_run(', which resulted in a training data size of ')
    p1.add_run(str(len(self.y_train)))
    p1.add_run(' and a testing data size of ')
    p1.add_run(str(len(self.y_test)))
    p1.add_run('. ')
    p1.add_run(self.report_model)
    if (self.lamNum != 0):
        p1.add_run(' was used as the for data analysis and its tuning parameters were selected using ')
        p1.add_run(self.report_tuning)
    else:
        p1.add_run(' was used as the for data analysis')
    p1.add_run(' with the training dataset.')
    
    document.add_heading('Data analysis summary', level=1)


    if self.model_name not in ['6', '13']:
        p2 = document.add_paragraph('The goodness-of-fit measured in adj R-squared is ')
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2.add_run(str(self.adj_rsquared))
        p2.add_run('Based on the best model selected, the prediction error on the testing data is ')
        p2.add_run(str(self.rmse_test_final))
        p2.add_run('. The residual plot is attached in the following section')
    else:
        p2 = document.add_paragraph('')
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2.add_run('Based on the best model selected, the prediction accuracy on the testing data is ')
        p2.add_run(str(self.testscore))

    
    document.add_heading('The performance measure table', level=1)

    if self.model_name not in ['6', '13']:
        performance_measure_table = [[1, "goodness-of-fit", self.adj_rsquared],
                                     [2, "training rmse", self.rmse_train_final]
            , [3, "testing rmse", self.rmse_test_final]]
    elif self.model_name == '13':
        performance_measure_table = [[1, "training accuracy", self.trainscore], [2, "testing accuracy", self.testscore]]
    else:
        performance_measure_table = [[1, "goodness-of-fit", self.adj_rsquared],
                                     [2, "training accuracy", self.trainscore]
            , [3, "training type I error", self.T1Error], [4, "training type II error", self.T2Error],
                                     [5, "testing accuracy", self.testscore],
                                     [6, "training type I error", self.T1Error1],
                                     [7, "training type II error", self.T2Error1]]
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Performance measure index'
    hdr_cells[1].text = 'Performance measure name'
    hdr_cells[2].text = 'Value'

    for item in performance_measure_table:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item[0])
        row_cells[1].text = str(item[1])
        row_cells[2].text = str(item[2])

    if self.model_name != '13':
        document.add_heading('The residual plot', level=1)
        if self.model_name == '7':
            for i in range(self.y_train.shape[1]):
                document.add_heading('task ' + str(i + 1), level=2)
                document.add_picture('./' + self.folderName + '/residual_plot' + str(i) + '.png', height=Inches(4),
                                     width=Inches(5.5))
        else:
            document.add_picture('./' + self.folderName + '/residual_plot.png', height=Inches(4), width=Inches(5.5))

    if (self.lamNum != 0):
        document.add_heading('The tuning plot', level=1)
        document.add_picture('./'+self.folderName+'/parameter_tuning_plot.png', height=Inches(4),width=Inches(5.5))
        
    document.add_page_break()
    document.save('./'+self.folderName+'/report_demo.docx')
